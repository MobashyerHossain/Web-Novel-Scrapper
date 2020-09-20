import requests
import os
from bs4 import BeautifulSoup as soup
from docx import Document
import regex as re

def nameCap(name):	
	nn = name.split("-")
	newName = ' '.join([str(part).capitalize() for part in nn])
	return newName

def pageValidity(link):
	validity = True
	pg = requests.get(link)
	pg_html = soup(pg.content, "html.parser")
	pg.close()

	cpt_link_list = []

	for cpt_clm in pg_html.find_all('ul', {'class' : 'list-chapter'}):
		cpt_links = cpt_clm.find_all('a')

		for cpt_link in cpt_links:
			cpt_link_list.append(cpt_link.get('href').replace("/omniscient-readers-viewpoint/", ""))	

	if len(cpt_link_list) < 50:
		validity = False

	return validity, cpt_link_list

def chapter_range(cpt_list):
	while True:
		st, ed = int(input("Starting Chapter : ")), str(input("Ending Chapter : "))

		# end chapter check
		if("last" in ed):
			ed = len(cpt_list)-1
		else:            
			for index, cpt in enumerate(cpt_list):
				if ed in cpt:
					ed = index
					break

		ed = int(ed)

		#start chapter check
		for index, cpt in enumerate(cpt_list):
			if str(st) in cpt:
				st = index
				break

		if st < ed:
			break
		else:
			print('Starting Chapter is Greater than Ending Chapter!!! Try Agian?')

	return st, ed

# Remove Unneccerary Staff
def rus(content):
	cpt = content
	# replace all unneccessary content
	# for Omniscient Reader
	cpt = cpt.replace("-ssi", "")
	cpt = cpt.replace("-nim", "")
	cpt = cpt.replace("Lv.", "level")
	cpt = cpt.replace("lv.", "level")
	cpt = cpt.replace("Ho-Seong", "Hosung")
	cpt = cpt.replace("Ho-seong", "Hosung")
	cpt = cpt.replace("Ha-Yeong", "Hayoung")
	cpt = cpt.replace("Ha-yeong", "Hayoung")
	cpt = cpt.replace("Jeong Hui-Won", "Jung Heewon")
	cpt = cpt.replace("Hui-Won", "Heewon")
	cpt = cpt.replace("Hui-won", "Heewon")
	cpt = cpt.replace("Yi Hyeon-Seong", "Lee Hyunsung")
	cpt = cpt.replace("Hyeon-Seong", "Hyunsung")
	cpt = cpt.replace("Hyeon-seong", "Hyunsung")
	cpt = cpt.replace("Yu Joong-Hyeok", "Yu Jungyeok")
	cpt = cpt.replace("Joong-Hyeok", "Jonghyuk")
	cpt = cpt.replace("Joong-hyeok", "Jonghyuk")
	cpt = cpt.replace("Jung-Hyeok", "Jonghyuk")
	cpt = cpt.replace("Jung-hyeok", "Jonghyuk")
	cpt = cpt.replace("Yi Gil-Yeong", "Lee Gilyoung")
	cpt = cpt.replace("Gil-Yeong", "Gilyoung")
	cpt = cpt.replace("Gil-yeong", "Gilyoung")
	cpt = cpt.replace("Shin Yu-Seung", "Shin Yoosung")
	cpt = cpt.replace("Yu-Seung", "Yoosung")
	cpt = cpt.replace("Yu-seung", "Yoosung")
	cpt = cpt.replace("Mister Dok-Ja", "Kim Dokja")
	cpt = cpt.replace("Dok-Ja", "Dokja")
	cpt = cpt.replace("Dok-ja", "Dokja")
	cpt = cpt.replace("Yu Sang-Ah", "Yoo Sangah")
	cpt = cpt.replace("Sang-Ah", "Sangah")
	cpt = cpt.replace("Sang-ah", "Sangah")
	cpt = cpt.replace("Han Myeong-Oh", "Han Myungoh")
	cpt = cpt.replace("Han Su-Yeong", "Han Suyeong")
	cpt = cpt.replace("Su-Yeong", "Suyeong")
	cpt = cpt.replace("Su-yeong", "Suyeong")
	cpt = cpt.replace("Myeong-Oh", "Myungoh")
	cpt = cpt.replace("Myeong-oh", "Myungoh")
	cpt = cpt.replace("Lightning Transformation", "Electrification")
	cpt = cpt.replace("Lamarck's Giraffe", "Lamarck's Kirin")
	cpt = cpt.replace("Yi Ji-Hye", "Lee Jihye")
	cpt = cpt.replace("Ji-Hye", "Jihye")
	cpt = cpt.replace("Ji-hye", "Jihye")
	cpt = cpt.replace("mantra", "True voice")
	cpt = cpt.replace("Lily Blooming on Aquarius", "Lily Pin of Aquarius")
	cpt = cpt.replace("cross-legged", "cross legged")
	cpt = cpt.replace("much-younger", "much younger")
	cpt = cpt.replace("Cheok Jun-Gyeong", "Cheok Jungyeong")
	cpt = cpt.replace("Jun-Gyeong", "Jungyeong")
	cpt = cpt.replace("Jun-gyeong", "Jungyeong")
	cpt = cpt.replace("Jeon Woo-Chi", "Jeon Woochi")
	cpt = cpt.replace("Woo-Chi", "Woochi")
	cpt = cpt.replace("Yi Seol-Hwa", "Lee Seolhwa")
	cpt = cpt.replace("Seol-Hwa", "Seolhwa")
	cpt = cpt.replace("Seol-hwa", "Seolhwa")
	cpt = cpt.replace("Yu-Shin", "Yushin")
	cpt = cpt.replace("Yu-shin", "Yushin")
	cpt = cpt.replace("Hak-Hyeon", "Hakhyeon")
	cpt = cpt.replace("Hak-hyeon", "Hakhyeon")
	cpt = cpt.replace("&lt;kim dok-ja=\"\" company=\"\"&gt;'s", "Kim Dokja's Company")
	cpt = cpt.replace("dok-ja", "dokja")    
	cpt = cpt.replace("Fable", "Story")
	cpt = cpt.replace("<", "'")
	cpt = cpt.replace(">", "'")

	# uncensoring censored words
	cpt = re.sub(r"[s|S][h|H|\*][i|I|\*][t|T]", "shit", cpt)
	cpt = re.sub(r"[f|F][u|U|\*][c|C|\*][k|K]", "fuck", cpt)
	cpt = re.sub(r"[b|B][i|i|\*][t|T|\*][c|C|\*][h|H]", "bitch", cpt)
	cpt = re.sub(r"[\*|a|A][s|S|\*][s|S|\*]", "ass", cpt)
	cpt = re.sub(r"[d|D][\*|a|A][\*|m|M][n|N]", "damn", cpt)
	cpt = re.sub(r"[b|B][\*|a|A][s|S|\*][t|T]\w*", "bastard", cpt)

	return cpt

def chapter(link):
	#get chapter page data
	cpt_pg = requests.get(link)
	cpt_pg_html = soup(cpt_pg.content, "html.parser")
	cpt_pg.close()

	cpt_title = cpt_pg_html.find("span", {"class":"chapter-text"}).get_text()

	cpt_content = cpt_pg_html.find(id = "chapter-content").find_all('p')

	cpt = ''

	for x in cpt_content:
		line = x.get_text()

		if line != "":
			cpt += "{}\n".format(line)

	return cpt_title, rus(cpt)


def main():
	LN_list = [
		"omniscient-readers-viewpoint",
	]

	#Options
	for index, opt in enumerate(LN_list, 1):
		print("{0:02d} : {nv}".format(index, nv = nameCap(opt)))
	print('\n')

	#start and end chapter index
	ln = -1
	while(ln < 0 or ln > len(LN_list) - 1):
		ln = int(input("Novel Number : ")) - 1

	#get index page data
	novel = LN_list[ln]
	idx_pg_link = "https://novelfull.com/{}.html".format(novel)
	idx_pg = requests.get(idx_pg_link)
	idx_pg_html = soup(idx_pg.content, "html.parser")
	idx_pg.close()

	# Create target Directory if don't exist
	dir_name = nameCap(novel)
	if not os.path.exists(dir_name):
		os.mkdir(dir_name)
		print("Directory {} is Created".format(dir_name))

	#find capter list
	cpt_list = []
	page_valid = True
	i = 0
	while(page_valid):
		i += 1
		pg_link = "https://novelfull.com/index.php/{}.html?page={}&amp;per-page=50".format(novel, i)
		page_valid, cpt_links = pageValidity(pg_link)
		cpt_list += cpt_links

	#check chapter validity
	st, ed = chapter_range(cpt_list)
	st_cpt = re.findall(r'\d+', cpt_list[st])[0]
	ed_cpt = re.findall(r'\d+', cpt_list[ed])[0]

	#document initialization
	LNdocument = Document()

	for index, cpt in enumerate(cpt_list[st:ed+1]):
		cpt_link = "https://novelfull.com/{}/{}".format(novel, cpt)
		cpt_title, cpt_content = chapter(cpt_link)

		cpt_idx = int(re.findall(r'\d+', cpt)[0])
		print("{0:03d} {title}".format(cpt_idx, title=cpt_title))

		#store chapter in doc file
		LNdocument.add_heading(cpt_title, level=1)
		LNdocument.add_paragraph(cpt_content)

		#devide every 50 chapter in different doc file
		if cpt_idx % 50 == 0:
			ed_cpt = re.findall(r'\d+', cpt_list[st+index])[0]
			doc_name = "{} {} - {}.docx".format(dir_name, st_cpt,ed_cpt)
			LNdocument.save("{}/{}".format(dir_name, doc_name))
			LNdocument = Document()
			st_cpt = str(int(ed_cpt)+1)

	#save final doc file
	ed_cpt = re.findall(r'\d+', cpt_list[ed])[0]
	doc_name = "{} {} - {}.docx".format(dir_name, st_cpt,ed_cpt)
	LNdocument.save("{}/{}".format(dir_name, doc_name))

if __name__=="__main__":
	main()