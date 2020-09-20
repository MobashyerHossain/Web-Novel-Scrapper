import requests
import os
from bs4 import BeautifulSoup as soup
from docx import Document
import regex as re

def chapter_range(cpt_list):
    while True:
        st, ed = int(input("Starting Chapter : ")), str(input("Ending Chapter : "))

        if("last" in ed):
        	ed = len(cpt_list)-1
        else:            
            for index, cpt in enumerate(cpt_list):
            	if ed in cpt.get_text():
                    ed = index
                    break

        ed = int(ed)

        #start chapter check
        for index, cpt in enumerate(cpt_list):
            if str(st) in cpt.get_text():
                st = index
                break

        if st < ed:
            break
        else:
            print('Starting Chapter is Greater than Ending Chapter!!! Try Agian?')

    return st, ed

def chapter(link):
    #get chapter page data
    cpt_pg = requests.get(link)
    cpt_pg_html = soup(cpt_pg.content, "html.parser")
    cpt_pg.close()

    cpt = cpt_pg_html.find(id="chaptercontent").prettify()

    #replace all unneccessary content
    cpt = cpt.replace("<br/>", "").replace("Translator: MrVoltaire1  ","").replace("Editor: Modlawls123","")
    cpt = cpt.replace("ChapterMid();","").replace("<script>","").replace("</script>","")
    cpt = cpt.replace('<div class="Readarea ReadAjax_content" id="chaptercontent">', "")
    cpt = cpt.replace('<div style="width:340;height:55px;overflow:hidden">', "")
    cpt = cpt.replace('<ins class="adsbygoogle" data-ad-client="ca-pub-2853920792116568" data-ad-slot="8807674870" style="display:inline-block;width:336px;height:50px">', "")
    cpt = cpt.replace('</ins>', "")
    cpt = cpt.replace('</div>', "")
    cpt = cpt.replace('Find authorized novels in Webnovel，faster updates, better experience，Please click www.webnovel.com  for visiting.', "")
    cpt = cpt.replace("(adsbygoogle = window.adsbygoogle || []).push({});", "")
    cpt = cpt.replace("Previous Chapter", "").replace("Next Chapter", "")    
    cpt = cpt.replace("Find authorized novels in Webnovel, faster updates, better experience, Please click for visiting.", "")

    # for Omniscient Reader
    cpt = cpt.replace("-ssi", "")
    cpt = cpt.replace("-nim", "")
    cpt = cpt.replace("Lv.", "level")
    cpt = cpt.replace("lv.", "level")
    cpt = cpt.replace("Ho-Seong", "Hosung")
    cpt = cpt.replace("Ho-seong", "Hosung")
    cpt = cpt.replace("Ha-Yeong", "Hayoung")
    cpt = cpt.replace("Ha-yeong", "Hayoung")
    cpt = cpt.replace("Jeong Hui-Won", "Jung Heewon")
    cpt = cpt.replace("Hui-Won", "Heewon")
    cpt = cpt.replace("Hui-won", "Heewon")
    cpt = cpt.replace("Yi Hyeon-Seong", "Lee Hyunsung")
    cpt = cpt.replace("Hyeon-Seong", "Hyunsung")
    cpt = cpt.replace("Hyeon-seong", "Hyunsung")
    cpt = cpt.replace("Yu Joong-Hyeok", "Yu Jungyeok")
    cpt = cpt.replace("Joong-Hyeok", "Jonghyuk")
    cpt = cpt.replace("Joong-hyeok", "Jonghyuk")
    cpt = cpt.replace("Jung-Hyeok", "Jonghyuk")
    cpt = cpt.replace("Jung-hyeok", "Jonghyuk")
    cpt = cpt.replace("Yi Gil-Yeong", "Lee Gilyoung")
    cpt = cpt.replace("Gil-Yeong", "Gilyoung")
    cpt = cpt.replace("Gil-yeong", "Gilyoung")
    cpt = cpt.replace("Shin Yu-Seung", "Shin Yoosung")
    cpt = cpt.replace("Yu-Seung", "Yoosung")
    cpt = cpt.replace("Yu-seung", "Yoosung")
    cpt = cpt.replace("Mister Dok-Ja", "Kim Dokja")
    cpt = cpt.replace("Dok-Ja", "Dokja")
    cpt = cpt.replace("Dok-ja", "Dokja")
    cpt = cpt.replace("Yu Sang-Ah", "Yoo Sangah")
    cpt = cpt.replace("Sang-Ah", "Sangah")
    cpt = cpt.replace("Sang-ah", "Sangah")
    cpt = cpt.replace("Han Myeong-Oh", "Han Myungoh")
    cpt = cpt.replace("Han Su-Yeong", "Han Suyeong")
    cpt = cpt.replace("Su-Yeong", "Suyeong")
    cpt = cpt.replace("Su-yeong", "Suyeong")
    cpt = cpt.replace("Myeong-Oh", "Myungoh")
    cpt = cpt.replace("Myeong-oh", "Myungoh")
    cpt = cpt.replace("Lightning Transformation", "Electrification")
    cpt = cpt.replace("Lamarck's Giraffe", "Lamarck's Kirin")
    cpt = cpt.replace("Yi Ji-Hye", "Lee Jihye")
    cpt = cpt.replace("Ji-Hye", "Jihye")
    cpt = cpt.replace("Ji-hye", "Jihye")
    cpt = cpt.replace("mantra", "True voice")
    cpt = cpt.replace("Lily Blooming on Aquarius", "Lily Pin of Aquarius")
    cpt = cpt.replace("cross-legged", "cross legged")
    cpt = cpt.replace("much-younger", "much younger")
    cpt = cpt.replace("Cheok Jun-Gyeong", "Cheok Jungyeong")
    cpt = cpt.replace("Jun-Gyeong", "Jungyeong")
    cpt = cpt.replace("Jun-gyeong", "Jungyeong")
    cpt = cpt.replace("Jeon Woo-Chi", "Jeon Woochi")
    cpt = cpt.replace("Woo-Chi", "Woochi")
    cpt = cpt.replace("Yi Seol-Hwa", "Lee Seolhwa")
    cpt = cpt.replace("Seol-Hwa", "Seolhwa")
    cpt = cpt.replace("Seol-hwa", "Seolhwa")
    cpt = cpt.replace("Yu-Shin", "Yushin")
    cpt = cpt.replace("Yu-shin", "Yushin")
    cpt = cpt.replace("Hak-Hyeon", "Hakhyeon")
    cpt = cpt.replace("Hak-hyeon", "Hakhyeon")
    cpt = cpt.replace("&lt;kim dok-ja=\"\" company=\"\"&gt;'s", "Kim Dokja's Company")
    cpt = cpt.replace("dok-ja", "dokja")    
    cpt = cpt.replace("Fable", "Story")

    #transform multiple emptylines to single empty line
    cpt = re.sub(r'\n\s*\n ', '\n\n', cpt)

    #uncensoring censored words
    cpt = cpt.replace("f*ck", "fuck").replace("sh*t", "shit").replace("*ss", "ass").replace("b*stard", "bastard")
    cpt = cpt.replace("b*tch", "bitch").replace("d*mn", "damn").replace("fu*k", "fuck")
    cpt = cpt.replace("F*ck", "Fuck").replace("Sh*t", "Shit").replace("*ss", "ass").replace("B*stard", "Bastard")
    cpt = cpt.replace("B*tch", "Bitch").replace("D*mn", "Damn").replace("Fu*k", "Fuck")

    return cpt

def main():
    LN_list = [
        ["Super Gene", "https://m.wuxiaworld.co/Super-Gene/"],
        ["Scholar's Advanced Technological System", "https://m.wuxiaworld.co/Scholar-s-Advanced-Technological-System/"],
    	["Reincarnation Of The Strongest Sword God", "https://m.wuxiaworld.co/Reincarnation-Of-The-Strongest-Sword-God/"],
    	["Realm of Myths and Legends", "https://m.wuxiaworld.co/Realm-of-Myths-and-Legends/"],
    	["The Legendary Mechanic", "https://m.wuxiaworld.co/The-Legendary-Mechanic/"],
    	["The Lord's Empire", "https://m.wuxiaworld.co/The-Lord%E2%80%99s-Empire/"],
    	["Monster Paradise", "https://m.wuxiaworld.co/Monster-Paradise/"],
    	["Imperial God Emperor", "https://m.wuxiaworld.co/Imperial-God-Emperor/"],
    	["Dragon Marked War God", "https://m.wuxiaworld.co/Dragon-Marked-War-God/"],
    	["Peerless Martial God", "https://m.wuxiaworld.co/Peerless-Martial-God/"],
    	["Against The Gods", "https://m.wuxiaworld.co/Against-the-Gods/"],
        ["Omniscient Reader", "https://m.wuxiaworld.co/Omniscient-Reader/"]
    ]

    #Options
    for index, opt in enumerate(LN_list, 1):
    	print("{0:0=2d}".format(index)+' : '+opt[0])

    #start and end chapter index
    ln = int(input("Novel Number : "))

    #get index page data
    idx_pg_link = LN_list[ln-1][1] + "all.html"
    idx_pg = requests.get(idx_pg_link)
    idx_pg_html = soup(idx_pg.content, "html.parser")
    idx_pg.close()

    # Create target Directory if don't exist
    dirName = idx_pg_html.header.span.get_text()
    if not os.path.exists(dirName):
        os.mkdir(dirName)
        print("Directory " , dirName ,  " Created ")

    #find chapter list
    cpt_list = idx_pg_html.find("div", {"id":"chapterlist"}).findAll("a", {"style":""})

    #check chapter validity
    st, ed = chapter_range(cpt_list)
    st_cpt = re.findall(r'\d+', cpt_list[st].get_text())[0]
    ed_cpt = re.findall(r'\d+', cpt_list[ed].get_text())[0]

    #document initialization
    LNdocument = Document()

    for index, cpt in enumerate(cpt_list[st:ed+1]):
        print(cpt.get_text())
        cpt_idx = int(re.findall(r'\d+', cpt.get_text())[0])
        #store chapter in doc file
        LNdocument.add_heading(cpt.get_text(), level=1)
        cpt_content = chapter(LN_list[ln-1][1] + cpt.get('href'))
        LNdocument.add_paragraph(cpt_content)

        #devide every 50 chapter in different doc file
        if cpt_idx % 50 == 0:
            ed_cpt = re.findall(r'\d+', cpt_list[st+index].get_text())[0]
            doc_name = idx_pg_html.header.span.get_text() + " {} - {}.docx".format(st_cpt,ed_cpt)
            LNdocument.save(dirName+'/'+doc_name)
            LNdocument = Document()
            st_cpt = str(int(ed_cpt)+1)

    #save final doc file
    ed_cpt = re.findall(r'\d+', cpt_list[ed].get_text())[0]
    doc_name = idx_pg_html.header.span.get_text() + " {} - {}.docx".format(st_cpt,ed_cpt)
    LNdocument.save(dirName+'/'+doc_name)

if __name__=="__main__":
    main()
